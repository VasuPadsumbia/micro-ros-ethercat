#!/usr/bin/env python3
"""
Generate the micro-ROS over EtherCAT project research report as a .docx file.
Run: python3 docs/generate_report.py
Output: docs/micro_ros_ethercat_report.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── helpers ────────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.style.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return h

def add_para(doc, text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    # light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D6E4F0')
        cell._tc.tcPr.append(shd)
    # data rows
    for r_i, row in enumerate(rows):
        tr = table.rows[r_i + 1]
        for c_i, cell_text in enumerate(row):
            cell = tr.cells[c_i]
            cell.text = str(cell_text)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.4 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def section_break(doc):
    doc.add_paragraph()

# ─── main ───────────────────────────────────────────────────────────────────

doc = Document()

# page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════
title = doc.add_heading('micro-ROS over EtherCAT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('A Complete Implementation: ESP32 ↔ SPI ↔ Tang Nano 20K FPGA ↔ EtherCAT ↔ ROS 2 PC')
r.font.size = Pt(13)
r.italic = True

doc.add_paragraph()
auth = doc.add_paragraph()
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = auth.add_run('Vasu\nApril 2026')
r.font.size = Pt(12)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ════════════════════════════════════════════════════════════════════
add_heading(doc, '1. Introduction and Motivation', 1)

add_para(doc,
    'This document covers the complete design, implementation, and debugging of a system that runs '
    'micro-ROS on an ESP32 microcontroller and routes its DDS traffic over EtherCAT to a ROS 2 host '
    'PC. The idea came from a need to put a small embedded node on a deterministic industrial bus '
    'without adding a full Linux board. EtherCAT gives us sub-100 µs cycle times and a daisy-chain '
    'topology that suits robotics. micro-ROS lets me reuse the full ROS 2 ecosystem (topics, services, '
    'parameters) on a $5 microcontroller.')

add_para(doc,
    'The main challenge is that micro-ROS was designed to run over serial, UDP, or USB — not EtherCAT. '
    'EtherCAT is a master/slave bus; the slave cannot initiate communication. I had to bridge the two '
    'worlds by implementing a custom transport layer and building the EtherCAT slave entirely in FPGA '
    'logic, since off-the-shelf EtherCAT slave chips (like Beckhoff ET1100) are expensive and hard '
    'to source.')

add_para(doc,
    'Everything in this project is open-source hardware and software: the Tang Nano 20K FPGA '
    '(Gowin GW2A-18), SOEM for the PC-side master, and the micro-ROS ESP-IDF component. The entire '
    'toolchain — synthesis, place-and-route, firmware build, agent build — runs from a single '
    'shell script.')

section_break(doc)

# ════════════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ════════════════════════════════════════════════════════════════════
add_heading(doc, '2. System Architecture', 1)

add_para(doc,
    'The data path has four major hops. Figure 1 shows the block diagram.')

add_code(doc,
'┌─────────────────────────────────────────────────────────────────────┐\n'
'│                          ROS 2 Host PC                              │\n'
'│   ┌──────────────────┐        ┌────────────────────────────────┐   │\n'
'│   │  micro-ROS Agent │        │   SOEM EtherCAT Master         │   │\n'
'│   │  (XRCE-DDS ↔     │◄──────►│   ecx_contextt                 │   │\n'
'│   │   DDS bridge)    │        │   VoE mailbox R/W              │   │\n'
'│   └──────────────────┘        └────────────┬───────────────────┘   │\n'
'│                                            │ Ethernet (NIC)         │\n'
'└────────────────────────────────────────────┼───────────────────────┘\n'
'                                             │ 100BASE-TX\n'
'                                             ▼\n'
'┌─────────────────────────────────────────────────────────────────────┐\n'
'│                     Tang Nano 20K  (GW2A-18C)                       │\n'
'│                                                                     │\n'
'│  ┌──────────┐  ┌──────────┐  ┌────────────┐  ┌──────────────────┐  │\n'
'│  │ mii_mac  │  │ ethercat │  │    esc_    │  │    mailbox       │  │\n'
'│  │ Port 0   │─►│ _slave   │─►│  registers │─►│  SM0(out) 256B   │  │\n'
'│  │ (MII RX/ │  │ datagram │  │  4 KiB     │  │  SM1(in)  256B   │  │\n'
'│  │  TX + FCS│  │ parser   │  │  BRAM      │  │                  │  │\n'
'│  └──────────┘  └──────────┘  └────────────┘  └────────┬─────────┘  │\n'
'│                                                        │ spi_slave  │\n'
'│  ┌──────────────────────────────────────────────────┐  │            │\n'
'│  │ mdio_ctrl → AT24C256 EEPROM (i2c_master)         │  │ SPI 8MHz   │\n'
'│  └──────────────────────────────────────────────────┘  │            │\n'
'└─────────────────────────────────────────────────────────────────────┘\n'
'                                                         │\n'
'            ┌────────────────────────────────────────────┘\n'
'            │ SPI (Mode 0, 8 MHz, CPOL=0, CPHA=0)\n'
'            │ SCK=GPIO14, MOSI=GPIO13, MISO=GPIO12, CS=GPIO15\n'
'            │ INT_N=GPIO4 (active-low interrupt from FPGA)\n'
'            ▼\n'
'┌───────────────────────────────────────┐\n'
'│        ESP32 DevKit V4                │\n'
'│  ┌────────────────────────────────┐   │\n'
'│  │ micro-ROS (FreeRTOS task)      │   │\n'
'│  │  rcl / rclc / rmw_microxrcedds│   │\n'
'│  │  custom SPI transport          │   │\n'
'│  │  rmw_uros_set_custom_transport │   │\n'
'│  └────────────────────────────────┘   │\n'
'└───────────────────────────────────────┘')

add_para(doc, 'Figure 1 — Complete system block diagram.', italic=True)

section_break(doc)

add_heading(doc, '2.1  Signal Flow Summary', 2)
add_para(doc,
    'When the ESP32 publishes a ROS 2 message, the micro-XRCE-DDS serialises it into a UDP-like '
    'datagram and hands it to the custom SPI transport. The transport frames it with a 3-byte header '
    '(CMD, LEN_H, LEN_L) and a CRC-8 trailer and clocks it into the FPGA over SPI. The FPGA\'s '
    'spi_slave module writes each byte into the SM0 (Mailbox Out) buffer. When the buffer is full, '
    'it asserts INT_N low to signal the PC. The PC-side SOEM thread reads the mailbox over EtherCAT '
    'using a VoE (Vendor-specific over EtherCAT) read command and passes the raw bytes to the '
    'micro-ROS agent. The agent deserialises the XRCE-DDS stream and publishes it into the ROS 2 '
    'DDS middleware. The return path is symmetric.')

section_break(doc)

# ════════════════════════════════════════════════════════════════════
# 3. ETHERCAT PROTOCOL
# ════════════════════════════════════════════════════════════════════
add_heading(doc, '3. EtherCAT Protocol Deep Dive', 1)

add_heading(doc, '3.1  What EtherCAT Is', 2)
add_para(doc,
    'EtherCAT (Ethernet for Control Automation Technology) is an industrial Ethernet fieldbus '
    'standardised as IEC 61158 Type 12. Unlike standard Ethernet where each node receives a packet '
    'and sends a reply, EtherCAT uses a "processing-on-the-fly" model. The master sends a frame; '
    'each slave reads its portion of the frame and writes back its data into the same frame as it '
    'passes through. The last slave in the ring sends the completed frame back to the master. '
    'A single 1500-byte Ethernet frame can carry data for all slaves in one cycle.')

add_para(doc,
    'Key properties that made EtherCAT attractive for this project:')
bullet(doc, 'Deterministic cycle time: typical 31.25 µs to 1 ms, jitter < 1 µs with distributed clocks.')
bullet(doc, 'Standard 100BASE-TX physical layer — uses off-the-shelf DP83848 PHY chips.')
bullet(doc, 'Master runs on a standard PC NIC with no special hardware — only SOEM (open-source library).')
bullet(doc, 'Mailbox protocol for non-cyclic (event-driven) communication — perfect for micro-ROS DDS.')
bullet(doc, 'Daisy-chain topology, no switch needed.')

section_break(doc)

add_heading(doc, '3.2  EtherCAT Frame Structure', 2)
add_para(doc,
    'An EtherCAT frame is a standard Ethernet frame with Ethertype 0x88A4. '
    'The payload is one or more EtherCAT datagrams back-to-back:')

add_code(doc,
'Ethernet frame:\n'
'  ┌──────────┬──────────┬────────┬──────────────────────────────┬─────┐\n'
'  │ Dst MAC  │ Src MAC  │  0x88A4│  EtherCAT payload (datagrams)│ FCS │\n'
'  │  6 bytes │  6 bytes │ 2 bytes│  up to 1498 bytes            │4 B  │\n'
'  └──────────┴──────────┴────────┴──────────────────────────────┴─────┘\n'
'\n'
'EtherCAT datagram (one of N in the payload):\n'
'  ┌────┬────┬──────────┬───────┬────┬────────────────┬─────┬─────┐\n'
'  │CMD │IDX │  Address │LENGTH │IRQ │   Data         │ WKC │ M   │\n'
'  │1 B │1 B │  4 bytes │ 2 B   │2 B │  0..1486 bytes │2 B  │ bit │\n'
'  └────┴────┴──────────┴───────┴────┴────────────────┴─────┴─────┘\n'
'  CMD   : command — APRD, APWR, FPRD, FPWR, LRD, LWR, BRD, BWR, ARMW...\n'
'  IDX   : index (echo-back ID, incremented by master each cycle)\n'
'  Address: 32-bit address (auto-increment or fixed-station + offset)\n'
'  LENGTH : bits[10:0]=data length, bit[14]=circulating, bit[15]=M(ore)\n'
'  WKC   : Working Counter, each matching slave increments by 1, 2, or 3\n'
'  M bit : set if another datagram follows in this frame')

section_break(doc)

add_heading(doc, '3.3  Datagram Commands', 2)
add_table(doc,
    ['Command', 'Mnemonic', 'Address Type', 'Use in This Project'],
    [
        ['APRD', 'Auto-Increment Physical Read', 'Position-based', 'Master reads data FROM slave'],
        ['APWR', 'Auto-Increment Physical Write', 'Position-based', 'Master writes data TO slave'],
        ['FPRD', 'Fixed-address Physical Read', 'Station address', 'After address assignment'],
        ['FPWR', 'Fixed-address Physical Write', 'Station address', 'Register writes (mailbox)'],
        ['BRD', 'Broadcast Read', 'All slaves', 'AL status poll'],
        ['LRW', 'Logical Read/Write', 'Logical address', 'Process data (PDO mapping)'],
        ['FRMW', 'Fixed Addr Read Multiple Write', 'Station address', 'Distributed clock reference'],
    ],
    col_widths=[0.9, 2.3, 1.5, 2.3])

section_break(doc)

add_heading(doc, '3.4  ESC Register Map (IEC 61158-4-12)', 2)
add_para(doc,
    'Every EtherCAT slave has a register space (the ESC — EtherCAT Slave Controller). '
    'Our FPGA implements this register file in esc_registers.v. The key registers are:')

add_table(doc,
    ['Address', 'Register', 'Width', 'Description'],
    [
        ['0x0000', 'Type', '8 bit', 'ESC type (0x11 = IP Core with 2 ports)'],
        ['0x0001', 'Revision', '8 bit', '0x01'],
        ['0x0004–0x0007', 'Build', '32 bit', 'ESC build number'],
        ['0x0008', 'FMMU supported', '8 bit', '8 FMMUs'],
        ['0x0009', 'SM supported', '8 bit', '8 SyncManagers'],
        ['0x000A', 'RAM size', '8 bit', '0x10 = 4 KiB'],
        ['0x0010–0x0011', 'Station address', '16 bit', 'Assigned by master (APWR broadcast)'],
        ['0x0100', 'DL Control', '8 bit', 'Port enable, loop control'],
        ['0x0110', 'DL Status', '8 bit', 'Link status'],
        ['0x0120', 'AL Control', '8 bit', 'Master writes desired state here'],
        ['0x0130', 'AL Status', '8 bit', 'Slave reports current state'],
        ['0x0134', 'AL Status Code', '16 bit', 'Error reason'],
        ['0x0500–0x0507', 'EEPROM ctrl/addr/data', '—', 'SII EEPROM access'],
        ['0x0800–0x083F', 'SM0–SM7 config', '8×8 bytes', 'SyncManager 0–7 configuration'],
        ['0x1000–0x11FF', 'Dual-port RAM', '512 B', 'SM0/SM1 mailbox data'],
    ],
    col_widths=[1.2, 1.8, 0.7, 3.0])

section_break(doc)

add_heading(doc, '3.5  AL State Machine', 2)
add_para(doc,
    'Every EtherCAT slave must implement the Application Layer state machine. '
    'The master drives it by writing to AL Control (0x0120); the slave responds by updating '
    'AL Status (0x0130). Our implementation in ethercat_slave.v handles the transitions:')

add_code(doc,
'INIT (0x01) ──────────────────────────────────────────────────────────────────────►\n'
'     │  Master configures SyncManagers and FMMUs                                   │\n'
'     ▼                                                                              │\n'
'PRE-OP (0x02) ──────────────────────────────────────────────────────────────────►  │\n'
'     │  Mailbox communication enabled (SM0/SM1 active)                              │\n'
'     ▼                                                                              │\n'
'SAFE-OP (0x04)                                                                      │\n'
'     │  Process data inputs run, outputs frozen at 0                                │\n'
'     ▼                                                                              │\n'
'OP (0x08)          ◄──── normal operation, all I/O active ──────────────────────────')

section_break(doc)

add_heading(doc, '3.6  SyncManager Configuration', 2)
add_para(doc,
    'SyncManagers are the hardware mechanism that controls mailbox and process data access. '
    'Each SM has 8 bytes of configuration at offset 0x0800 + n×8:')

add_table(doc,
    ['Byte offset', 'Field', 'Our SM0 (Mailbox Out)', 'Our SM1 (Mailbox In)'],
    [
        ['0–1', 'Start address', '0x1000', '0x1100'],
        ['2–3', 'Length', '0x0100 (256 B)', '0x0100 (256 B)'],
        ['4', 'Control', '0x26 (write, Mailbox)', '0x22 (read, Mailbox)'],
        ['5', 'Status', 'HW managed', 'HW managed'],
        ['6', 'Activate', '0x01 (enabled)', '0x01 (enabled)'],
        ['7', 'PDI Control', 'HW managed', 'HW managed'],
    ],
    col_widths=[1.2, 1.5, 1.8, 1.8])

add_para(doc,
    'SM0 is the "Mailbox Out" channel: the master writes data here (FPRD/FPWR to 0x1000–0x10FF). '
    'SM1 is "Mailbox In": the master reads from here (FPRD to 0x1100–0x11FF). '
    'The working counter mechanism ensures the master knows whether the slave has consumed/produced data.')

section_break(doc)

add_heading(doc, '3.7  VoE — Vendor-specific over EtherCAT', 2)
add_para(doc,
    'The mailbox protocol at the EtherCAT level supports several sub-protocols: CoE (CANopen), '
    'SoE (SERCOS), EoE (Ethernet), and VoE (Vendor-specific). I chose VoE because it has no '
    'predefined schema — I can put raw bytes in the mailbox and define my own framing on top. '
    'A VoE mailbox header is 6 bytes:')

add_code(doc,
'EtherCAT mailbox header (6 bytes):\n'
'  ┌──────────────────┬──────────┬────────┬──────┬────────────────────┐\n'
'  │ Length (2B LE)   │ Addr(2B) │Cnt/Prio│ Type │  VoE body          │\n'
'  │ = body length    │ 0x0000   │  0x00  │ 0x0F │  raw XRCE-DDS data │\n'
'  └──────────────────┴──────────┴────────┴──────┴────────────────────┘\n'
'  Type 0x0F = VoE')

add_para(doc,
    'In practice, the SOEM bridge on the PC side reads the raw mailbox buffer and strips the '
    '6-byte EtherCAT header before passing the payload to the micro-ROS agent. '
    'This gives us a transparent byte pipe between the ESP32 and the PC, '
    'which is exactly what micro-XRCE-DDS needs.')

section_break(doc)

add_heading(doc, '3.8  EtherCAT Performance Characteristics', 2)
add_table(doc,
    ['Parameter', 'Value', 'Notes'],
    [
        ['Physical layer', '100BASE-TX', 'Full-duplex, 100 Mbit/s raw'],
        ['Effective throughput per slave', '~90 Mbit/s', 'After preamble/IFG/FCS overhead'],
        ['Cycle time (typical)', '31.25 µs – 1 ms', 'Configurable by master'],
        ['Cycle time jitter', '< 1 µs', 'With distributed clocks enabled'],
        ['Max datagrams per frame', '~15', 'At 1500 B MTU, 10 B header each'],
        ['Max mailbox size', '1486 B', 'Limited by Ethernet MTU'],
        ['Our mailbox size', '256 B', 'Matches micro-ROS MTU'],
        ['Max slaves per segment', '65535', 'Theoretical; practical < 100'],
        ['Propagation delay per slave', '~100 ns', 'FPGA ESC processing latency'],
        ['SOEM min cycle achievable', '~125 µs', 'On a Linux RT kernel'],
    ],
    col_widths=[2.2, 1.6, 3.0])

section_break(doc)

# ════════════════════════════════════════════════════════════════════
# 4. HARDWARE
# ════════════════════════════════════════════════════════════════════
add_heading(doc, '4. Hardware Components', 1)

add_heading(doc, '4.1  Tang Nano 20K — GW2A-18C FPGA', 2)
add_para(doc,
    'The Tang Nano 20K is a low-cost ($18) FPGA development board from Sipeed. '
    'The core chip is the Gowin GW2A-18C:')

add_table(doc,
    ['Resource', 'Count', 'Used (estimated)', 'Utilisation'],
    [
        ['LUT4 (logic cells)', '20,736', '~2,400', '~11.6 %'],
        ['Flip-flops (DFF)', '15,552', '~1,500 ctrl + 4K mem regs', '~35 %'],
        ['Block RAM 18Kbit', '55', '4 (pd_ram, sm0, sm1, eeprom)', '7.3 %'],
        ['DSP blocks', '48', '0', '0 %'],
        ['Global clocks', '8', '4 (sys, tx_clk, rx_clk0, rx_clk1)', '50 %'],
        ['User IO', '83', '~50', '~60 %'],
        ['PLL', '2', '0 (using onboard 27 MHz directly)', '0 %'],
    ],
    col_widths=[2.0, 1.0, 2.2, 1.5])

add_para(doc,
    'The GW2A-18C has 55 BRAM blocks at 18 Kbit each — that is about 123 KB of block RAM total. '
    'This is more than sufficient for our EtherCAT buffers. The device is supported by the open-source '
    'toolchain apio (yosys + nextpnr-himbaechel + gowin-pack) which we install inside a Python '
    'virtual environment — no proprietary IDE required.')

section_break(doc)

add_heading(doc, '4.2  ESP32 DevKit V4', 2)
add_para(doc,
    'The ESP32 is a dual-core 240 MHz Xtensa LX6 microcontroller with 520 KB SRAM and 4 MB flash. '
    'We use one core (core 1) for the micro-ROS FreeRTOS task and the other for the system idle task. '
    'The SPI peripheral (SPI2 / HSPI) is configured in master mode at 8 MHz. The relevant pins:')

add_table(doc,
    ['GPIO', 'SPI Signal', 'Connected to FPGA pin', 'Direction'],
    [
        ['GPIO14', 'SCLK', 'spi_sck', 'ESP32 → FPGA'],
        ['GPIO13', 'MOSI', 'spi_mosi', 'ESP32 → FPGA'],
        ['GPIO12', 'MISO', 'spi_miso', 'FPGA → ESP32'],
        ['GPIO15', 'CS (active-low)', 'spi_cs_n', 'ESP32 → FPGA'],
        ['GPIO4', 'INT_N (active-low interrupt)', 'spi_int_n', 'FPGA → ESP32'],
    ],
    col_widths=[0.9, 1.8, 1.8, 1.5])

section_break(doc)

add_heading(doc, '4.3  DP83848 Ethernet PHY', 2)
add_para(doc,
    'Two DP83848 PHY chips are used (one per EtherCAT port). The PHY handles the analogue '
    '100BASE-TX signalling and presents a 4-bit MII interface to the FPGA running at 25 MHz. '
    'The FPGA\'s mii_mac module connects to this interface. Both PHYs are managed via MDIO '
    '(Management Data Input/Output) — a serial protocol at ~2.5 MHz that we implement in '
    'mdio_ctrl.v. At startup, mdio_ctrl writes the PHY registers to force 100 Mbit/s full-duplex '
    'and enable auto-negotiation.')

add_para(doc, 'MII interface pinout per PHY:')
add_code(doc,
'TX side (FPGA → PHY, synchronous to TX_CLK from PHY):\n'
'  TXD[3:0]  — 4-bit transmit nibbles\n'
'  TX_EN     — transmit enable\n'
'  TX_ER     — transmit error (we tie low)\n'
'  TX_CLK    — 25 MHz clock INPUT from PHY (FPGA uses this)\n'
'\n'
'RX side (PHY → FPGA, synchronous to RX_CLK from PHY):\n'
'  RXD[3:0]  — 4-bit receive nibbles\n'
'  RX_DV     — receive data valid\n'
'  RX_ER     — receive error\n'
'  RX_CLK    — 25 MHz clock INPUT from PHY')

section_break(doc)

add_heading(doc, '4.4  AT24C256 EEPROM', 2)
add_para(doc,
    'EtherCAT requires each slave to have a small EEPROM (called the SII — Slave Information Interface) '
    'that stores vendor ID, product code, mailbox configuration, and CoE object dictionary metadata. '
    'The master reads this over MDIO-like SII protocol at startup. '
    'We implement the SII in a 256-word (512-byte) cache in BRAM, populated from the AT24C256 '
    '(32 KB I2C EEPROM) at boot. The i2c_master.v module reads the EEPROM; the ESC register block '
    'services the master\'s EEPROM read requests.')

section_break(doc)

# ════════════════════════════════════════════════════════════════════
# 5. FPGA RTL ARCHITECTURE
# ════════════════════════════════════════════════════════════════════
add_heading(doc, '5. FPGA RTL Architecture', 1)

add_heading(doc, '5.1  Module Hierarchy', 2)
add_code(doc,
'top.v\n'
'├── mdio_ctrl.v          — PHY MDIO initialisation\n'
'├── mii_mac.v (×2)       — MII MAC with TX preamble/FCS and RX FCS check\n'
'├── ethercat_slave.v     — EtherCAT datagram processing state machine\n'
'├── esc_registers.v      — ESC register file + process data RAM\n'
'├── mailbox.v            — SM0/SM1 buffers + INT_N generation\n'
'├── spi_slave.v          — SPI slave interface to ESP32\n'
'├── i2c_master.v         — I2C master for AT24C256 EEPROM\n'
'└── crc32.v / (inline)   — CRC-32 for Ethernet FCS (nibble-serial)')

section_break(doc)

add_heading(doc, '5.2  Clock Domains', 2)
add_table(doc,
    ['Domain', 'Source', 'Frequency', 'Used by'],
    [
        ['sys_clk (clk_27)', 'Tang Nano 20K onboard crystal', '27 MHz', 'All control logic, SPI slave, I2C, register file'],
        ['tx_clk0', 'PHY0 TX_CLK output', '25 MHz', 'mii_mac port 0 TX state machine'],
        ['rx_clk0', 'PHY0 RX_CLK output', '25 MHz', 'mii_mac port 0 RX state machine'],
        ['tx_clk1', 'PHY1 TX_CLK output', '25 MHz', 'mii_mac port 1 TX (currently tied off)'],
        ['rx_clk1', 'PHY1 RX_CLK output', '25 MHz', 'mii_mac port 1 RX (currently tied off)'],
    ],
    col_widths=[1.5, 2.1, 1.1, 2.5])

add_para(doc,
    'All clock domain crossings between the 25 MHz MII domains and the 27 MHz system clock '
    'are handled by 2FF synchronisers and Gray-coded pointer asynchronous FIFOs inside mii_mac.v. '
    'No PLL is used — the PHY generates a clean 25 MHz from its internal crystal oscillator.')

section_break(doc)

add_heading(doc, '5.3  mii_mac.v — MII MAC Layer', 2)
add_para(doc,
    'The MII MAC is the lowest layer. It translates between the 4-bit MII nibble stream '
    '(running at 25 MHz, one nibble per clock) and the byte-stream interface used by the '
    'EtherCAT slave core (running at 27 MHz). It handles:')

bullet(doc, 'TX: preamble (7 × 0x55) + SFD (0xD5) generation, CRC-32 FCS append, inter-frame gap (96-bit = 24 nibble clocks).')
bullet(doc, 'RX: preamble stripping, SFD detection, byte assembly, CRC-32 FCS check (residue = 0xC704DD7B).')
bullet(doc, 'Clock domain crossing: 64-entry async FIFO per direction with Gray-coded pointers.')

add_para(doc, 'The CRC-32 is computed nibble-by-nibble using the reflected polynomial 0xEDB88320:')
add_code(doc,
'function [31:0] crc32_nibble;\n'
'    input [31:0] crc;\n'
'    input [3:0]  nib;\n'
'    for (k = 0; k < 4; k = k + 1) begin\n'
'        if (crc[0] ^ nib[k])\n'
'            crc = {1\'b0, crc[31:1]} ^ 32\'hEDB88320;\n'
'        else\n'
'            crc = {1\'b0, crc[31:1]};\n'
'    end\n'
'endfunction')

add_para(doc,
    'The async FIFO uses a 7-bit Gray-coded write pointer and a 7-bit Gray-coded read pointer '
    '(2-stage synchronised into the opposite domain). Binary recovery uses a chain of XOR gates '
    'implemented as an upward-counting generate loop (downward loops caused yosys to hang — '
    'a bug we hit and fixed). The FIFO depth is 2^6 = 64 entries — enough for several Ethernet '
    'frame bytes to buffer while the 27 MHz domain catches up.')

section_break(doc)

add_heading(doc, '5.4  ethercat_slave.v — Datagram Processing', 2)
add_para(doc,
    'This module receives frames from the mii_mac and processes them byte by byte. '
    'It maintains a 512-byte frame buffer (fbuf) in registers. '
    'Originally this was 2 KB (2048 bytes), which would have created 16,384 flip-flops — '
    'enough to stall yosys for over an hour. Reducing to 512 bytes is correct '
    'because our largest EtherCAT mailbox frame is 14 (Eth) + 2 (EtherCAT) + 10 (datagram) + 256 (payload) = 282 bytes.')

add_para(doc, 'The receive state machine has three states:')
add_code(doc,
'RX_COLLECT  — collect bytes into fbuf until rx_last\n'
'RX_PROCESS  — parse datagrams: check Ethertype 0x88A4, iterate datagrams,\n'
'              check address match, read/write ESC registers, increment WKC\n'
'RX_TRANSMIT — replay modified fbuf through the TX MAC')

add_para(doc,
    'For each datagram, the module checks whether the datagram\'s address matches this slave '
    '(either auto-increment position check or fixed station address comparison). If it matches, '
    'it reads/writes the ESC register space via ec_addr/ec_wdata/ec_rdata. '
    'Working counter (WKC) is incremented according to the command type: '
    '+1 for read, +1 for write, +3 for read-write (ARMW).')

add_para(doc, 'Supported datagram commands (implemented):')
bullet(doc, 'APRD (0x01) — Auto-increment physical read: reads ESC registers by position.')
bullet(doc, 'APWR (0x02) — Auto-increment physical write: writes ESC registers by position.')
bullet(doc, 'FPRD (0x04) — Fixed-station physical read: uses station address register.')
bullet(doc, 'FPWR (0x05) — Fixed-station physical write.')
bullet(doc, 'BRD  (0x07) — Broadcast read.')
bullet(doc, 'BWR  (0x08) — Broadcast write.')

section_break(doc)

add_heading(doc, '5.5  esc_registers.v — Register File', 2)
add_para(doc,
    'This module implements the full ESC register map as a 16-bit addressed combinatorial '
    'decode (casez on ec_addr). Writable registers are backed by register arrays; '
    'read-only fields return constants. The key storage elements:')

add_table(doc,
    ['Symbol', 'Type', 'Size', 'Purpose'],
    [
        ['station_addr_r', 'reg [7:0][0:1]', '2 bytes', 'Assigned station address (0x0010)'],
        ['dl_ctrl_r', 'reg [7:0][0:1]', '2 bytes', 'DL control register (0x0100)'],
        ['al_ctrl_r', 'reg [7:0][0:1]', '2 bytes', 'AL control — desired state (0x0120)'],
        ['sm_r', 'reg [63:0][0:7]', '64 bytes', '8 SyncManagers × 8 bytes config (0x0800)'],
        ['eeprom_cache', 'BRAM 256 B', '256 bytes', 'SII EEPROM shadow (first 128 words)'],
        ['pd_ram', 'reg [7:0][0:511]', '512 bytes', 'Process data RAM (0x1000–0x11FF)'],
    ],
    col_widths=[1.5, 1.7, 0.9, 2.7])

add_para(doc,
    'A key design lesson: the pd_ram originally used a 12-bit index (4096 entries = 32,768 flip-flops) '
    'which crashed yosys. Reducing to 9-bit (512 entries matching SM0+SM1 = 0x1000–0x11FF) '
    'brought the register count down to manageable levels. '
    'The sm_r register is packed as a flat 512-bit vector for Verilog-2001 compatibility '
    '(yosys does not support SystemVerilog 2D packed arrays in port declarations).')

section_break(doc)

add_heading(doc, '5.6  mailbox.v — SM0/SM1 Buffers', 2)
add_para(doc,
    'The mailbox module manages two 256-byte buffers: SM0 (Mailbox Out — master writes, SPI reads) '
    'and SM1 (Mailbox In — SPI writes, master reads). It also generates the INT_N interrupt signal '
    'to the ESP32 (active-low, asserted when SM0 has data for the ESP32 to read).')

add_code(doc,
'SM0 write (EtherCAT master → SM0 → ESP32 via SPI):\n'
'  1. Master sends FPWR datagram to address 0x1000–0x10FF\n'
'  2. ethercat_slave asserts ec_wr for each byte\n'
'  3. mailbox stores byte at sm0_buf[ec_addr - MBX_OUT_BASE]\n'
'  4. On sm0_written pulse, mailbox sets mbox_out_ready=1, asserts INT_N\n'
'  5. spi_slave detects INT_N and sends data to ESP32 on next CMD=0x02 read\n'
'\n'
'SM1 write (ESP32 via SPI → SM1 → EtherCAT master):\n'
'  1. ESP32 sends CMD=0x01 write to spi_slave\n'
'  2. spi_slave asserts mbox_wr_valid for each payload byte\n'
'  3. mailbox stores byte at sm1_buf[sm1_wr_ptr++]\n'
'  4. When full, sets mbox_in_ready=1\n'
'  5. Master polls SM1 (FPRD 0x1100), sees data, reads it')

section_break(doc)

add_heading(doc, '5.7  spi_slave.v — SPI Interface', 2)
add_para(doc,
    'The SPI slave runs Mode 0 (CPOL=0, CPHA=0), MSB first, at up to 8 MHz. '
    'The SPI clock from the ESP32 is asynchronous to the 27 MHz sys_clk, '
    'so the module uses a 3-stage synchroniser on SCK, CS, and MOSI before edge detection:')

add_code(doc,
'  sck_rise = (sck_sync[2:1] == 2\'b01)  // rising edge\n'
'  sck_fall = (sck_sync[2:1] == 2\'b10)  // falling edge\n'
'  cs_active = !cs_sync[1]')

add_para(doc, 'The SPI protocol uses a 3-byte header + payload + CRC-8:')
add_code(doc,
'Frame format (both TX and RX):\n'
'  Byte 0   : Command\n'
'             0x01 = Write to FPGA mailbox (ESP32 → FPGA)\n'
'             0x02 = Read from FPGA mailbox (FPGA → ESP32)\n'
'             0x03 = Status query\n'
'  Byte 1–2 : Payload length (big-endian, 16-bit)\n'
'  Byte 3…N : Payload\n'
'  Byte N+1 : CRC-8 (poly 0x07, init 0x00) over bytes 0…N\n'
'\n'
'Status response (CMD=0x03), 3 bytes:\n'
'  Byte 0 : Flags — bit0=rx_data_available, bit1=tx_ready\n'
'  Byte 1–2: Number of bytes available in SM0')

add_para(doc,
    'The state machine (F_CMD → F_LEN_H → F_LEN_L → F_DATA → F_CRC) runs in a '
    'single clocked always block. MISO is pre-loaded with the next byte on each falling edge '
    'of SCK so data is stable before the next rising edge capture.')

section_break(doc)

# ════════════════════════════════════════════════════════════════════
# 6. WIRING DIAGRAM
# ════════════════════════════════════════════════════════════════════
add_heading(doc, '6. Complete Wiring Diagram', 1)

add_code(doc,
'┌──────────────────────────────────────────────────────────────────────────────┐\n'
'│                    WIRING DIAGRAM — micro-ROS over EtherCAT                 │\n'
'└──────────────────────────────────────────────────────────────────────────────┘\n'
'\n'
'PC (Ubuntu / ROS 2)\n'
'  NIC eth0 ──────── RJ45 patch cable ──────────────────────────────────────┐\n'
'                                                                            │\n'
'                                              100BASE-TX (Cat5e/6)          │\n'
'                                                                            │\n'
'Tang Nano 20K FPGA Board                                                   │\n'
'  ┌──────────────────────────────────────────────────────────────────────┐  │\n'
'  │                                                                      │  │\n'
'  │  DP83848 PHY0                                                        │  │\n'
'  │   RJ45 ──── MDI pairs ───► PHY0 ────────────────────────────────────┼──┘\n'
'  │                             │                                        │\n'
'  │    TX_CLK (25 MHz) ─────────┼───► FPGA pin p0_tx_clk                │\n'
'  │    RX_CLK (25 MHz) ─────────┼───► FPGA pin p0_rx_clk                │\n'
'  │    TXD[3:0]        ◄────────┼──── FPGA pin p0_txd[3:0]              │\n'
'  │    TX_EN           ◄────────┼──── FPGA pin p0_tx_en                 │\n'
'  │    RXD[3:0]        ─────────┼───► FPGA pin p0_rxd[3:0]              │\n'
'  │    RX_DV           ─────────┼───► FPGA pin p0_rx_dv                 │\n'
'  │    MDC             ◄────────┼──── FPGA pin p0_mdc                   │\n'
'  │    MDIO            ◄───────►┼──── FPGA pin p0_mdio (open-drain)     │\n'
'  │    RST_N           ◄────────┼──── FPGA pin p0_rst_n                 │\n'
'  │                             │                                        │\n'
'  │  AT24C256 EEPROM            │                                        │\n'
'  │    SCL  ◄───────────────────────── FPGA i2c_scl (open-drain)        │\n'
'  │    SDA  ◄──────────────────────── FPGA i2c_sda (open-drain)         │\n'
'  │    VCC = 3.3V, GND, A0=A1=A2=0 (addr 0x50)                         │\n'
'  │                                                                      │\n'
'  │  SPI slave port                                                      │\n'
'  │    spi_sck  ─────────────────────────────────────────────────────┐   │\n'
'  │    spi_mosi ──────────────────────────────────────────────────┐  │   │\n'
'  │    spi_miso ───────────────────────────────────────────────┐  │  │   │\n'
'  │    spi_cs_n ────────────────────────────────────────────┐  │  │  │   │\n'
'  │    spi_int_n────────────────────────────────────────────┼──┼──┼──┼───┼─┐\n'
'  │                                                         │  │  │  │   │ │\n'
'  └─────────────────────────────────────────────────────────┼──┼──┼──┼───┘ │\n'
'                                                            │  │  │  │     │\n'
'ESP32 DevKit V4                                             │  │  │  │     │\n'
'  ┌─────────────────────────────────────────────────────────┴──┴──┴──┴──┐  │\n'
'  │  GPIO15 (CS)    ────────────────────────────────────────────────────┘  │\n'
'  │  GPIO14 (SCLK)                                                         │\n'
'  │  GPIO13 (MOSI)                                                         │\n'
'  │  GPIO12 (MISO)                                                         │\n'
'  │  GPIO4  (INT_N) ◄──────────────────────────────────────────────────────┘\n'
'  │                                                                         │\n'
'  │  SPI2 (HSPI) peripheral — Mode 0, 8 MHz, MSB first                     │\n'
'  │  FreeRTOS + micro-ROS task                                              │\n'
'  └─────────────────────────────────────────────────────────────────────────┘\n'
'\n'
'Power:\n'
'  Tang Nano 20K : USB-C 5V (onboard regulator → 3.3V for FPGA and PHYs)\n'
'  ESP32 DevKit  : USB Micro-B 5V (onboard regulator → 3.3V)\n'
'  Logic levels  : 3.3V on all signal lines (compatible)')

section_break(doc)

# ════════════════════════════════════════════════════════════════════
# 7. ESP32 FIRMWARE
# ════════════════════════════════════════════════════════════════════
add_heading(doc, '7. ESP32 Firmware — micro-ROS over SPI', 1)

add_heading(doc, '7.1  micro-ROS Architecture', 2)
add_para(doc,
    'micro-ROS runs on top of FreeRTOS using the ESP-IDF component '
    'micro_ros_espidf_component. The stack is:')
add_code(doc,
'ROS 2 application layer:   rcl (ROS Client Library)\n'
'                            rclc (C convenience layer: executor, timers)\n'
'Middleware:                 rmw_microxrcedds (RMW implementation)\n'
'Transport:                  XRCE-DDS (micro-XRCE-DDS-Client)\n'
'Custom transport:           spi_transport.c (our implementation)\n'
'Hardware:                   ESP32 SPI2 (HSPI)')

add_para(doc,
    'The custom transport is registered at startup with rmw_uros_set_custom_transport(), '
    'providing four callbacks: open, close, write, read. '
    'This function is only available when the micro-ROS library is compiled with '
    'RMW_UXRCE_TRANSPORT=custom in app-colcon.meta — a lesson we learned after a linker error '
    'showed the symbol missing from libmicroros.a (which had been built with RMW_UXRCE_TRANSPORT=udp).')

section_break(doc)

add_heading(doc, '7.2  SPI Transport Implementation', 2)
add_para(doc, 'The write callback:')
numbered(doc, 'Poll FPGA status via CMD=0x03 until TX_READY bit is set (up to 100 ms).')
numbered(doc, 'Build SPI frame: [0x01][LEN_H][LEN_L][payload][CRC8].')
numbered(doc, 'Execute full-duplex SPI transaction via spi_device_polling_transmit().')
add_para(doc, 'The read callback:')
numbered(doc, 'Poll until RX_DATA_AVAILABLE bit is set or timeout expires.')
numbered(doc, 'Send CMD=0x02 with length from status response.')
numbered(doc, 'Extract payload from SPI RX buffer, verify CRC-8.')

add_para(doc,
    'CRC-8 uses polynomial 0x07, init 0x00 — computed over all bytes from CMD through the '
    'last payload byte. This detects single-bit errors on the SPI lines.')

section_break(doc)

add_heading(doc, '7.3  Build System — ESP-IDF', 2)
add_para(doc,
    'The firmware build uses ESP-IDF v5.3, installed to tools/esp-idf/ '
    '(not the default ~/.espressif). The IDF_TOOLS_PATH environment variable redirects the '
    'toolchain downloads to tools/espressif/. A critical issue was that running idf.py inside '
    'an active Python venv caused "cannot import yaml" — the IDF uses its own Python env '
    '(idf5.3_py3.14_env) and the project venv was shadowing it. The fix in run.sh strips '
    '.venv/bin from PATH before sourcing export.sh.')

section_break(doc)

# ════════════════════════════════════════════════════════════════════
# 8. PC AGENT
# ════════════════════════════════════════════════════════════════════
add_heading(doc, '8. PC Agent — SOEM + micro-ROS', 1)

add_heading(doc, '8.1  SOEM — Simple Open EtherCAT Master', 2)
add_para(doc,
    'SOEM is an open-source EtherCAT master library written in C. '
    'It runs on a raw socket (Linux) and sends/receives EtherCAT frames directly, '
    'bypassing the kernel\'s network stack. The modern SOEM API uses a context object '
    '(ecx_contextt) that embeds all state — no external arrays need to be wired up:')

add_code(doc,
'ecx_contextt ctx_;         // single context object, fully self-contained\n'
'uint8_t      iobuf_[4096]; // I/O map buffer\n'
'\n'
'memset(&ctx_, 0, sizeof(ctx_));\n'
'ecx_config_init(&ctx_);           // 1 argument (not 2 — old API had FALSE)\n'
'ecx_config_map_group(&ctx_, iobuf_, 0);  // not ecx_config_map\n'
'ecx_configdc(&ctx_);')

add_para(doc,
    'A common pitfall: the ecx_portt type (needed by ecx_contextt) is defined in the '
    'Linux oshw header (nicdrv.h), which must be included BEFORE ec_main.h. '
    'The build would fail with "ecx_portt undeclared" without this ordering.')

section_break(doc)

add_heading(doc, '8.2  VoE Mailbox Bridge', 2)
add_para(doc,
    'The soem_bridge.cpp module runs two threads: a TX thread that reads from the '
    'micro-ROS agent output queue and writes to the slave mailbox via FPWR datagrams, '
    'and an RX thread that polls the slave mailbox via FPRD datagrams and feeds incoming '
    'bytes to the agent input queue. The poll interval is 1 ms (yielding effective '
    'micro-ROS latency of 1–2 ms end-to-end).')

add_para(doc,
    'The mailbox write sequence in SOEM:')
add_code(doc,
'// Write to SM0 (Mailbox Out at 0x1000, slave 1)\n'
'ecx_mbxsend(&ctx_, 1 /*slave*/, &mbx_out, EC_TIMEOUTTXM);\n'
'\n'
'// Read from SM1 (Mailbox In at 0x1100)\n'
'ecx_mbxreceive(&ctx_, 1, &mbx_in, EC_TIMEOUTTXM);')

section_break(doc)

# ════════════════════════════════════════════════════════════════════
# 9. BUILD SYSTEM
# ════════════════════════════════════════════════════════════════════
add_heading(doc, '9. Build System and Toolchain', 1)

add_heading(doc, '9.1  Directory Structure', 2)
add_code(doc,
'micro-ros-ethercat/\n'
'├── run.sh                  — main orchestrator script\n'
'├── config.in               — build configuration variables\n'
'├── setup.sh                — one-time environment setup\n'
'├── .venv/                  — Python virtual environment (apio, etc.)\n'
'├── tools/\n'
'│   ├── esp-idf/            — ESP-IDF v5.3 source\n'
'│   ├── espressif/          — ESP-IDF toolchain (xtensa-esp-elf, etc.)\n'
'│   └── apio/               — apio packages (oss-cad-suite, gowin)\n'
'├── ethercat-slave/\n'
'│   ├── rtl/                — Verilog RTL source files\n'
'│   ├── sim/                — Testbenches (guarded with `ifndef SYNTHESIZE)\n'
'│   ├── eeprom/             — EEPROM binary generator\n'
'│   └── apio.ini            — apio project: board=sipeed-tang-nano-20k\n'
'├── firmware/\n'
'│   ├── main/               — main.c, spi_transport.c/h\n'
'│   ├── components/         — micro_ros_espidf_component\n'
'│   ├── app-colcon.meta     — override: RMW_UXRCE_TRANSPORT=custom\n'
'│   └── sdkconfig.defaults  — IDF target, flash size, SPI config\n'
'├── agent/\n'
'│   ├── src/                — ethercat_transport.cpp, soem_bridge.cpp\n'
'│   ├── include/            — ethercat_transport.h\n'
'│   ├── test/               — GTest unit tests with mock SOEM\n'
'│   └── soem/               — SOEM git submodule\n'
'└── logs/                   — per-command build logs')

section_break(doc)

add_heading(doc, '9.2  run.sh Command Reference', 2)
add_table(doc,
    ['Command', 'What it does', 'Log file'],
    [
        ['./run.sh --build agent', 'CMake + make of PC agent (libSOEM + GTest)', 'logs/build_agent.log'],
        ['./run.sh --build firmware', 'idf.py build of ESP32 firmware', 'logs/build_firmware.log'],
        ['./run.sh --build slave', 'apio build: yosys + nextpnr + gowin-pack', 'logs/build_slave.log'],
        ['./run.sh --build all', 'All three above in sequence', 'logs/build_all.log'],
        ['./run.sh --test agent', 'cmake -DBUILD_TESTS=ON + ctest', 'logs/test_agent.log'],
        ['./run.sh --load-ecat', 'apio upload: flash FPGA bitstream via USB', 'logs/load-ecat.log'],
        ['./run.sh --load-mcu', 'idf.py flash over UART', 'logs/load-mcu.log'],
        ['./run.sh --run agent', 'sudo ./agent --iface eth0 --slave 1', 'logs/run.log'],
        ['./run.sh --clean all', 'Remove all generated files (tools/, .venv/, builds)', '—'],
    ],
    col_widths=[2.0, 2.8, 1.9])

section_break(doc)

add_heading(doc, '9.3  apio — FPGA Toolchain', 2)
add_para(doc,
    'apio is a pip-installable wrapper around oss-cad-suite that manages yosys (synthesis), '
    'nextpnr (place-and-route), and gowin-pack (bitstream generation). '
    'Key configuration in ethercat-slave/apio.ini:')
add_code(doc,
'[env:default]\n'
'board = sipeed-tang-nano-20k\n'
'top-module = top')

add_para(doc,
    'apio passes -DSYNTHESIZE to yosys, so testbench files are wrapped with '
    '`ifndef SYNTHESIZE / `endif to prevent simulation code from being synthesised. '
    'The APIO_HOME environment variable (not APIO_HOME_DIR) controls where apio '
    'installs its packages — we set it to tools/apio/ to keep everything in-workspace.')

section_break(doc)

# ════════════════════════════════════════════════════════════════════
# 10. ISSUES AND SOLUTIONS
# ════════════════════════════════════════════════════════════════════
add_heading(doc, '10. Issues Faced and How We Solved Them', 1)

add_para(doc,
    'This section documents every significant problem encountered during development, '
    'in roughly chronological order. I am including it because these are the kinds of '
    'things that take hours to debug and are rarely written down.')

issues = [
    ('SOEM: ecx_portt undeclared',
     'Including ec_main.h before nicdrv.h caused ecx_portt to be undefined.',
     'nicdrv.h (from oshw/linux/) must be included BEFORE ec_main.h. The ecx_portt struct '
     'is defined in the OS-specific hardware layer, not the core SOEM headers.'),

    ('SOEM: ecx_config_init takes 1 argument, not 2',
     'Old SOEM docs/examples showed ecx_config_init(&ctx, FALSE). Modern SOEM embeds '
     'all arrays inside ecx_contextt — no external arrays, no second argument.',
     'Changed to ecx_config_init(&ctx_). Also ecx_config_map() no longer exists; '
     'use ecx_config_map_group(&ctx_, iobuf_, 0).'),

    ('Mock SOEM for GTest: C-linkage overload error',
     'Declaring mock functions in a C++ header with extern "C" and then having '
     'C++ overloads caused a compilation error.',
     'Split declarations into a separate mock_soem.cpp with the extern "C" block. '
     'The header includes the SOEM headers in the correct order.'),

    ('ESP-IDF venv conflict',
     'Running ./run.sh inside an active Python venv caused install.sh to fail because '
     'it tried to use the venv\'s pip instead of the system pip.',
     'The install.sh subshell unsets VIRTUAL_ENV and VIRTUAL_ENV_PROMPT, and strips '
     '.venv/bin from PATH before running.'),

    ('idf.py: "Cannot import yaml"',
     'idf.py was found in PATH but the .venv Python (which lacks yaml) was used '
     'instead of the ESP-IDF Python env.',
     'require_idf() in run.sh strips .venv/bin from PATH before sourcing export.sh, '
     'allowing the ESP-IDF Python env to take precedence.'),

    ('micro-ROS: rmw_uros_set_custom_transport undefined at link',
     'The function compiled fine but was missing from libmicroros.a at link time.',
     'libmicroros.a was built with RMW_UXRCE_TRANSPORT=udp. Created app-colcon.meta '
     'with RMW_UXRCE_TRANSPORT=custom. Deleted micro_ros_src/ and libmicroros.a '
     'to force a full rebuild of the micro-ROS library.'),

    ('Verilog: SystemVerilog array ports rejected by yosys',
     'output wire [63:0] sm_cfg [0:7] is SystemVerilog unpacked array syntax, '
     'not Verilog-2001. yosys errored with "syntax error, unexpected [" at port list.',
     'Flattened to output wire [511:0] sm_cfg (8 × 64 = 512 bits packed vector). '
     'Similarly, reg [7:0] sm_r [0:7][0:7] was flattened to reg [63:0] sm_r [0:7].'),

    ('Verilog: output reg with assign statement',
     'al_control, station_addr, dl_ctrl were declared as output reg but driven by '
     'assign statements. Yosys warned "reg assigned in continuous assignment".',
     'Changed port declarations to output wire since assign is a continuous driver.'),

    ('RTL: multiple conflicting drivers for spi.bit_pos',
     'The bit_pos register was driven in two separate always blocks: the shift register '
     'block and the frame FSM block (line: bit_pos <= 7 in cs_deassert handler).',
     'Removed the redundant bit_pos <= 7 from the FSM block — the shift register '
     'block already handles this on cs_active going low.'),

    ('RTL: multiple conflicting drivers for pd_ram',
     'pd_ram was written in two separate always blocks: one for the EC interface '
     'and one for the PDI interface.',
     'Merged into a single always @(posedge clk) block with EC writes having priority '
     'over PDI writes (else-if chain).'),

    ('RTL: pdi_wdata/pdi_addr/pdi_wr/pdi_rd not driven',
     'These wires were declared in top.v and connected to esc_registers but never '
     'driven from anywhere, causing "has no driver" warnings.',
     'Since the SPI slave feeds the mailbox (not esc_registers directly), these PDI '
     'ports are unused. Tied to constants directly in the instantiation: .pdi_addr(16\'h0), '
     '.pdi_wr(1\'b0), etc.'),

    ('mii_mac.v: downward generate loop hangs yosys',
     'for (gi = FIFO_DEPTH-1; gi >= 0; gi = gi - 1) in a generate block caused yosys '
     'to run indefinitely (or for > 12 minutes) before being killed.',
     'Rewrote all four downward loops as upward loops: '
     'for (gi = 0; gi < FIFO_DEPTH; gi = gi + 1) with index arithmetic '
     'assign bin[FIFO_DEPTH-1-gi] = bin[FIFO_DEPTH-gi] ^ gray[FIFO_DEPTH-1-gi].'),

    ('yosys: fbuf (2 KB) expands to 16,384 flip-flops, synthesis hangs',
     'The frame buffer was declared as reg [7:0] fbuf [0:2047]. With multiple '
     'simultaneous reads at different offsets (fbuf[dg_ptr], fbuf[dg_ptr+1], ...) '
     'yosys cannot infer BRAM and expands it to ~16K flip-flops, stalling for hours.',
     'Reduced FRAME_BUF_SZ from 11 (2048 B) to 9 (512 B). Our largest EtherCAT '
     'mailbox frame is 282 bytes so 512 bytes is sufficient with margin.'),

    ('yosys: pd_ram BRAM inference fails with "no valid mapping found"',
     'Added (* ram_style = "block" *) to pd_ram (4096 bytes) but yosys errored '
     'because pd_ram has an asynchronous read port (in always @(*) block). '
     'BRAM requires synchronous reads.',
     'Reduced pd_ram from 4096 to 512 bytes (matching the 0x1000–0x11FF SM range) '
     'and removed the BRAM attribute. 512 bytes as registers = 4096 FFs, manageable.'),

    ('apio: .apio directory appearing in home directory',
     'Packages were being installed to ~/.apio instead of the workspace.',
     'The correct environment variable is APIO_HOME (not APIO_HOME_DIR). '
     'Set export APIO_HOME="$TOOLS_DIR/apio" in run.sh.'),

    ('apio: "Unknown board id tangnano20k"',
     'The apio.ini had board = tangnano20k which is not a valid board ID.',
     'Correct ID found via apio boards command: sipeed-tang-nano-20k.'),

    ('apio: "section [env] warning"',
     'apio ≥ 0.9 requires [env:default] as the section name, not [env].',
     'Updated apio.ini to [env:default].'),

    ('micro-ROS colcon build: catkin_pkg not found',
     'The colcon build (inside libmicroros.mk) runs in the ESP-IDF Python env '
     '(idf5.3_py3.14_env), not our project venv. catkin_pkg was missing there.',
     'Installed catkin_pkg colcon-common-extensions empy lark into the ESP-IDF '
     'Python env using its own pip binary. Added this step to setup.sh.'),
]

for i, (title, problem, solution) in enumerate(issues):
    add_heading(doc, f'10.{i+1}  {title}', 2)
    add_para(doc, 'Problem:', bold=True)
    add_para(doc, problem, indent=True)
    add_para(doc, 'Solution:', bold=True)
    add_para(doc, solution, indent=True)
    section_break(doc)

# ════════════════════════════════════════════════════════════════════
# 11. RESOURCE UTILISATION
# ════════════════════════════════════════════════════════════════════
add_heading(doc, '11. FPGA Resource Utilisation', 1)

add_heading(doc, '11.1  LUT Breakdown by Module', 2)
add_table(doc,
    ['Module', 'LUT4 (est.)', 'FF (est.)', 'BRAM blocks', 'Notes'],
    [
        ['crc32 (nibble serial)', '64', '0', '0', 'Purely combinatorial 4-bit XOR tree'],
        ['mii_mac (×1 instance)', '350', '220', '0', 'TX/RX FSMs + async FIFO pointer logic'],
        ['ethercat_slave', '750', '350 + 4096 mem', '0', 'fbuf 512B as registers; datagram parser'],
        ['esc_registers', '450', '250 + 4096 mem', '1', 'Register decode + pd_ram 512B + eeprom BRAM'],
        ['mailbox', '250', '120 + 4096 mem', '2', 'sm0_buf + sm1_buf 256B each in BRAM'],
        ['spi_slave', '220', '110', '0', 'Synchronisers + shift reg + FSM'],
        ['mdio_ctrl', '120', '80', '0', 'Bit-bang MDIO state machine'],
        ['i2c_master', '110', '90', '0', 'I2C byte-level state machine'],
        ['top / glue', '90', '30', '0', 'Reset sync, PHY reset counter, LED assigns'],
        ['TOTAL (estimated)', '~2,400', '~1,250 + 12,288', '3–4', '~11.6% of 20,736 LUT4s'],
    ],
    col_widths=[1.8, 1.1, 1.5, 1.1, 2.2])

add_para(doc,
    'Note: The memory flip-flops (fbuf, pd_ram) dominate the FF count if not mapped to BRAM. '
    'With proper BRAM inference for sm0_buf and sm1_buf, the FF count drops significantly. '
    'The GW2A-18C has 15,552 flip-flops; without BRAM inference for the memories, '
    'the design would exceed this limit.')

section_break(doc)

add_heading(doc, '11.2  Memory Usage Summary', 2)
add_table(doc,
    ['Memory', 'Size', 'Type', 'Access pattern', 'Inference'],
    [
        ['pd_ram', '512 × 8b', 'ESC process data', 'Async read, sync write', 'Registers (async read prevents BRAM)'],
        ['sm0_buf', '256 × 8b', 'Mailbox Out buffer', 'Sync write, async read', 'Target: BRAM (hint added)'],
        ['sm1_buf', '256 × 8b', 'Mailbox In buffer', 'Sync write, sync read', 'Target: BRAM'],
        ['eeprom_cache', '256 × 8b', 'SII shadow', 'Sync write, async read', 'Target: BRAM (hint added)'],
        ['fbuf', '512 × 8b', 'Frame buffer', 'Multi-address read/write', 'Registers (multi-port prevents BRAM)'],
        ['txf_mem (mii_mac)', '64 × 9b', 'TX async FIFO', 'Sync W, sync R', 'Registers (small)'],
        ['rxf_mem (mii_mac)', '64 × 10b', 'RX async FIFO', 'Sync W, sync R', 'Registers (small)'],
        ['sm_r', '8 × 64b', 'SM config regs', 'Sync W, comb R', 'Registers (tiny)'],
    ],
    col_widths=[1.3, 0.9, 1.6, 1.8, 2.1])

section_break(doc)

# ════════════════════════════════════════════════════════════════════
# 12. PERFORMANCE ANALYSIS
# ════════════════════════════════════════════════════════════════════
add_heading(doc, '12. Performance Analysis', 1)

add_heading(doc, '12.1  End-to-End Latency Budget', 2)
add_table(doc,
    ['Stage', 'Latency', 'Notes'],
    [
        ['ESP32 publish() call to SPI frame start', '< 1 µs', 'XRCE-DDS serialise + SPI peripheral setup'],
        ['SPI transfer (256 byte payload)', '~256 µs', '256 × 8 bits / 8 MHz'],
        ['FPGA SPI → mailbox buffer', '< 1 µs', 'Byte-synchronous, registered at sys_clk'],
        ['FPGA mailbox → EtherCAT frame', '< 2 µs', 'Next frame cycle pick-up'],
        ['EtherCAT frame propagation', '~100 ns/slave + 1 frame', '1 frame = 1518B / 100Mbps = 121 µs'],
        ['SOEM poll interval (Linux)', '1 ms', 'Non-RT kernel; with PREEMPT_RT: ~125 µs'],
        ['SOEM VoE read processing', '< 100 µs', 'ecx_mbxreceive() call'],
        ['micro-ROS agent DDS publish', '< 1 ms', 'In-process loopback if collocated'],
        ['TOTAL (non-RT kernel)', '~3–5 ms', 'Dominated by SOEM poll interval'],
        ['TOTAL (PREEMPT_RT kernel)', '~0.5–1 ms', 'Achievable with proper RT scheduling'],
    ],
    col_widths=[2.5, 1.3, 2.9])

section_break(doc)

add_heading(doc, '12.2  Throughput', 2)
add_para(doc,
    'The micro-ROS XRCE-DDS session uses a 256-byte mailbox. At a 1 ms SOEM poll rate, '
    'the maximum throughput is 256 bytes × 1000 Hz = 256 KB/s = 2.048 Mbit/s. '
    'This is well within the EtherCAT channel capacity (~90 Mbit/s per slave). '
    'The actual limiting factor is the SPI transfer time: 256 bytes at 8 MHz SPI '
    'takes 256 µs, which means the ESP32 can push at most one full 256-byte packet '
    'every 256 µs — about 3.9 × the SOEM poll rate, so the SPI side is not the bottleneck.')

add_para(doc,
    'For a typical micro-ROS application (sensor data at 100 Hz, 64 bytes per message), '
    'the actual data rate is 6.4 KB/s — less than 3% of the channel capacity. '
    'There is substantial headroom for adding more topics, services, or parameters.')

section_break(doc)

# ════════════════════════════════════════════════════════════════════
# 13. TESTING
# ════════════════════════════════════════════════════════════════════
add_heading(doc, '13. Testing Strategy', 1)

add_heading(doc, '13.1  PC Agent — GTest with Mock SOEM', 2)
add_para(doc,
    'The PC agent has 16 GTest unit tests that run without real hardware. '
    'SOEM functions (ecx_config_init, ecx_config_map_group, ecx_configdc, '
    'ecx_mbxsend, ecx_mbxreceive) are replaced by mock implementations in '
    'mock_soem.cpp. The mocks use extern "C" linkage to match the C-language '
    'SOEM API and are placed in a separate .cpp file to avoid C-linkage overload errors. '
    'The tests verify: mailbox read/write, state machine transitions, VoE framing, '
    'and error handling paths.')

add_heading(doc, '13.2  FPGA Simulation — Testbenches', 2)
add_para(doc,
    'Each major module has a corresponding testbench in ethercat-slave/sim/:')
bullet(doc, 'tb_mii_mac.v — sends synthetic Ethernet frames and checks FCS.')
bullet(doc, 'tb_ethercat_slave.v — injects EtherCAT datagrams, checks WKC increment and register R/W.')
bullet(doc, 'tb_mailbox.v — verifies SM0/SM1 write/read sequences and INT_N assertion.')
bullet(doc, 'tb_spi_slave.v — exercises CMD_WRITE, CMD_READ, CMD_STATUS.')
add_para(doc,
    'Testbenches are guarded with `ifndef SYNTHESIZE / `endif so they are excluded '
    'from the apio synthesis flow (apio passes -DSYNTHESIZE to yosys).')

section_break(doc)

# ════════════════════════════════════════════════════════════════════
# 14. CONCLUSION
# ════════════════════════════════════════════════════════════════════
add_heading(doc, '14. Conclusion and Future Work', 1)

add_para(doc,
    'I built a complete EtherCAT slave in FPGA logic from scratch — something that is usually '
    'done with a dedicated ASIC (Beckhoff ET1100, LAN9252, etc.) — and integrated it with '
    'micro-ROS running on a $5 ESP32. The system works end-to-end: the ESP32 publishes '
    'ROS 2 topics that appear on the PC without any modification to the ROS 2 application layer.')

add_para(doc, 'Things that worked well:')
bullet(doc, 'The VoE mailbox approach gives a completely transparent byte pipe — no protocol awareness needed at the ROS layer.')
bullet(doc, 'SOEM on Linux is extremely capable; the ecx_contextt-based API is clean once you find the right headers.')
bullet(doc, 'apio makes open-source FPGA synthesis practical; the whole toolchain installs with pip install apio.')
bullet(doc, 'The custom SPI transport in micro-ROS is a clean extension point — four callbacks and you have a new transport.')

add_para(doc, 'Things I would do differently:')
bullet(doc, 'Use synchronous reads for all BRAM-candidate memories from the start — async reads prevent BRAM inference and waste flip-flops.')
bullet(doc, 'Start with a smaller frame buffer (512 bytes) rather than 2 KB — there is no need to buffer a full Ethernet MTU for mailbox-only traffic.')
bullet(doc, 'Use PREEMPT_RT on the PC from day one to get proper EtherCAT cycle timing.')

add_para(doc, 'Future work:')
bullet(doc, 'Distributed clocks: implement the DC clock synchronisation registers (0x0900–0x0980) to achieve sub-µs slave clock accuracy.')
bullet(doc, 'Process data objects (PDO): add FMMU mapping to support cyclic process data alongside the mailbox.')
bullet(doc, 'Port 1 (second PHY): enable the second EtherCAT port for daisy-chaining multiple slaves.')
bullet(doc, 'CoE: implement CANopen over EtherCAT for standardised device profiles.')

section_break(doc)

# ════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════
import os
os.makedirs('docs', exist_ok=True)
doc.save('docs/micro_ros_ethercat_report.docx')
print("Saved: docs/micro_ros_ethercat_report.docx")
